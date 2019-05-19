from docx import Document
from docx.shared import Cm, Pt
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

def makeTable(document):
    table = document.add_table(rows=10, cols=4)
    table.style = "TableGrid"

    # table size
    widthList = [3, 7, 3, 3]
    for row in table.rows:
        row.height = Cm(0.8)
        for col, width in zip(row.cells, widthList):
            col.width = Cm(width)

    # merge cells
    title = table.cell(0, 0)
    titleRun = title.paragraphs[0].add_run("Title: ")
    titleRun.font.size = Pt(20)
    titleRun.bold = True

    titleContentStart = table.cell(0, 1)
    titleContentEnd = table.cell(0, 3)
    titleContent = titleContentStart.merge(titleContentEnd)
    titleContentRun = titleContent.paragraphs[0].add_run("이곳에 제목이 들어갑니다.")

    content = table.cell(1, 0)
    contentRun = content.paragraphs[0].add_run("Content: ")
    contentRun.font.size = Pt(20)
    contentRun.bold = True

    contentStart = table.cell(1, 1)
    contentEnd = table.cell(8, 3)
    content = contentStart.merge(contentEnd)
    contentRun = content.paragraphs[0].add_run("이곳에 내용이 들어갑니다.")

    nums = table.cell(9, 2)
    numsRun = nums.paragraphs[0].add_run("nums: ")
    numsRun.font.size = Pt(20)
    numsRun.bold = True

    numCount = table.cell(9, 3)
    numCountRun = numCount.paragraphs[0].add_run("현재 글자수/최대 글자수")



    # align center
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return document

if __name__=="__main__":
    document = Document()
    for i in range(2):
        document = makeTable(document)
        document.add_page_break()
    document.save('{}_{}_{}.docx'.format(datetime.now().hour, datetime.now().minute, datetime.now().second))