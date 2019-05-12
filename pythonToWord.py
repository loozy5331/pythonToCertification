from docx import Document
from docx.shared import Cm, Pt
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

document = Document()

table = document.add_table(rows=20, cols=8)
table.style = "TableGrid"

# table size
widthList = [0.7, 0.7, 0.7, 2, 2, 2, 2, 2]
for row in table.rows:
    row.height = Cm(0.8)
    for col, width in zip(row.cells, widthList):
        col.width = Cm(width*1.27)

# merge cells
pictureStart = table.cell(0, 0)
pictureEnd = table.cell(4, 2)
picture = pictureStart.merge(pictureEnd)
pictureRun = picture.paragraphs[0].add_run("사진")
pictureRun.font.size = Pt(12)
pictureRun.bold = True

titleStart = table.cell(0, 3)
titleEnd = table.cell(1, 7)
title = titleStart.merge(titleEnd)
titleRun = title.paragraphs[0].add_run("이력서")
titleRun.font.size = Pt(20)
titleRun.bold = True

nameStart = table.cell(2, 3)
nameEnd = table.cell(3, 3)
name = nameStart.merge(nameEnd)
nameRun = name.paragraphs[0].add_run("이름")
nameRun.bold =True

nameContentStart = table.cell(2, 4)
nameContentEnd = table.cell(3, 5)
nameContent = nameContentStart.merge(nameContentEnd)

sex = table.cell(2, 6)
sexRun = sex.paragraphs[0].add_run("성별")
sexRun.bold =True

age = table.cell(2, 7)
ageRun = age.paragraphs[0].add_run("나이")
ageRun.bold =True

addressStart = table.cell(5, 0)
addressEnd = table.cell(5, 2)
address = addressStart.merge(addressEnd)
addressRun = address.paragraphs[0].add_run("주소")
addressRun.bold =True

addressContentStart = table.cell(5, 3)
addressContentEnd = table.cell(5, 7)
addressContent = addressContentStart.merge(addressContentEnd)

emailStart = table.cell(6, 0)
emailEnd = table.cell(6, 2)
email = emailStart.merge(emailEnd)
emailRun = email.paragraphs[0].add_run("e-mail")
emailRun.bold =True

highNameStart = table.cell(7, 0)
highNameEnd = table.cell(7, 2)
highName = highNameStart.merge(highNameEnd)
highNameRun = highName.paragraphs[0].add_run("고등학교 이름")
highNameRun.bold =True

highNameContent = table.cell(7, 3)

highAdDate = table.cell(7, 4)
highAdDateRun = highAdDate.paragraphs[0].add_run("입학일")
highAdDateRun.bold =True

highAdDateContent = table.cell(7, 5)

highGrDate = table.cell(7, 6)
highGrDateRun = highGrDate.paragraphs[0].add_run("졸업일")
highGrDateRun.bold =True

highGrDateContent = table.cell(7, 7)

univNameStart = table.cell(8, 0)
univNameEnd = table.cell(8, 2)
univName = univNameStart.merge(univNameEnd)
univNameRun = univName.paragraphs[0].add_run("대학교 이름")
univNameRun.bold =True

univNameContent = table.cell(8, 3)

univAdDate = table.cell(8, 4)
univAdDateRun = univAdDate.paragraphs[0].add_run("입학일")
univAdDateRun.bold =True

univAdDateContent = table.cell(8, 5)

univGrDate = table.cell(8, 6)
univGrDateRun = univGrDate.paragraphs[0].add_run("졸업일")
univGrDateRun.bold =True

compTitleStart = table.cell(9, 0)
compTitleEnd = table.cell(9, 7)
compTitle = compTitleStart.merge(compTitleEnd)
compTitleRun = compTitle.paragraphs[0].add_run("경력사항")
compTitleRun.font.size = Pt(15)
compTitleRun.bold =True

compDateStart = table.cell(10, 0)
compDateEnd = table.cell(10, 2)
compDate = compDateStart.merge(compDateEnd)
compDateRun = compDate.paragraphs[0].add_run("근무기간(연/월/일)")
compDateRun.bold =True

compNameStart = table.cell(10, 3)
compNameEnd = table.cell(10, 6)
compName = compNameStart.merge(compNameEnd)
compNameRun = compName.paragraphs[0].add_run("회사명")
compNameRun.bold =True

compPlace = table.cell(10, 7)
compPlaceRun = compPlace.paragraphs[0].add_run("근무지")
compPlaceRun.bold =True

compDateContentStart = table.cell(11, 0)
compDateContentEnd = table.cell(11, 2)
compDateContent = compDateContentStart.merge(compDateContentEnd)

compNameContentStart = table.cell(11, 3)
compNameContentEnd = table.cell(11, 6)
compNameContent = compNameContentStart.merge(compNameContentEnd)

compPlaceContent = table.cell(11, 7)

certDateStart = table.cell(14, 0)
certDateEnd = table.cell(14, 2)
certDate = certDateStart.merge(certDateEnd)
certDateRun = certDate.paragraphs[0].add_run("취득일(연/월/일)")
certDateRun.bold =True

certNameStart = table.cell(14, 3)
certNameEnd = table.cell(14, 6)
certName = certNameStart.merge(certNameEnd)
certNameRun = certName.paragraphs[0].add_run("자격증")
certNameRun.bold =True

certPlace = table.cell(14, 7)
certPlaceRun = certPlace.paragraphs[0].add_run("비고")
certPlaceRun.bold =True


certDateContentStart = table.cell(15, 0)
certDateContentEnd = table.cell(15, 2)
certDateContent = certDateContentStart.merge(certDateContentEnd)

certNameContentStart = table.cell(15, 3)
certNameContentEnd = table.cell(15, 6)
certNameContent = certNameContentStart.merge(certNameContentEnd)

certPlaceContent = table.cell(15, 7)

# align center
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

document.save('{}_{}_{}.docx'.format(datetime.now().hour, datetime.now().minute, datetime.now().second))